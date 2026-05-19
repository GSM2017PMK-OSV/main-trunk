import os
import re
import sys
import xml.etree.ElementTree as ET
from datetime import datetime

from docx import Document


def extract_formulas(doc):
    """Извлекает математические формулы из документа"""
    formulas = []
    namespace = {
        'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math',
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }

    for p in doc.paragraphs:
        xml_str = p._element.xml
        if 'm:oMath' in xml_str:
            try:
                root = ET.fromstring(xml_str)
                for omath in root.findall('.//m:oMath', namespace):
                    formula_text = extract_math_text(omath, namespace)
                    if formula_text:
                        formulas.append(formula_text)
            except Exception as e:
                formulas.append(f"Ошибка извлечения формулы: {str(e)}")

    # Дополнительный поиск текстовых формул
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text])
    text_formulas = re.findall(
    r'(?:Формула|Уравнение)\s*[:\-]\s*(.+?)(?=\n|$)',
    full_text,
     re.IGNORECASE)
    formulas.extend(text_formulas)

    return formulas


def extract_math_text(element, ns):
    """Рекурсивно извлекает текст из математических элементов"""
    text = ""
    for child in element:
        if child.tag.endswith('}r'):
            for t in child.findall('.//w:t', ns):
                text += t.text if t.text else ''
        elif child.tag.endswith('}f'):
            for frac in child.findall('.//m:num', ns):
                text += extract_math_text(frac, ns) + '/'
            for frac in child.findall('.//m:den', ns):
                text += extract_math_text(frac, ns)
        elif child.tag.endswith('}lim'):
            text += "lim_"
            for low in child.findall('.//m:limLow', ns):
                for e in low.findall('.//m:e', ns):
                    text += extract_math_text(e, ns)
        elif child.tag.endswith('}int'):
            text += "∫"
            for sub in child.findall('.//m:sub', ns):
                text += extract_math_text(sub, ns)
            for sup in child.findall('.//m:sup', ns):
                text += extract_math_text(sup, ns)
        elif child.tag.endswith('}rad'):
            text += "√"
            for deg in child.findall('.//m:deg', ns):
                text += extract_math_text(deg, ns) + "("
            for e in child.findall('.//m:e', ns):
                text += extract_math_text(e, ns) + ")"
        elif child.tag.endswith('}sup'):
            text += "^"
            for e in child.findall('.//m:e', ns):
                text += extract_math_text(e, ns)
        elif child.tag.endswith('}sub'):
            text += "_"
            for e in child.findall('.//m:e', ns):
                text += extract_math_text(e, ns)
        elif child.tag.endswith('}e'):
            text += extract_math_text(child, ns)
        elif child.tag.endswith('}d'):
            text += extract_math_text(child, ns)
        elif child.tag.endswith('}func'):
            text += extract_math_text(child, ns) + "("
        elif child.tag.endswith('}fName'):
            text += extract_math_text(child, ns)
        elif child.tag.endswith('}argPr'):
            text += extract_math_text(child, ns) + ")"
        else:
            text += extract_math_text(child, ns)

    return text


def analyze_document(file_path):
    """Проводит полный анализ документа"""
    try:
        doc = Document(file_path)
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text])

        # Извлечение формул
        formulas = extract_formulas(doc)

        # Анализ разделов
        sections = analyze_sections(full_text)

        # Анализ требований
        requirements = analyze_requirements(full_text, formulas)

        # Анализ методов верификации
        verification = analyze_verification(full_text)

        # Формирование отчета
        report = {
            "metadata": {
                "document": os.path.basename(file_path),
                "analysis_date": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "pages": len(doc.paragraphs) // 50 + 1
            },
            "formulas": formulas,
            "sections": sections,
            "requirements": requirements,
            "verification": verification,
            "recommendations": generate_recommendations(sections, requirements, verification, formulas)
        }

        # Сохранение отчета
        report_path = save_report(report, file_path)
        return report_path

    except Exception as e:
        return f"Ошибка анализа: {str(e)}"


def analyze_sections(text):
    """Анализирует наличие ключевых разделов"""
    sections = {
        "Теоретические положения": bool(re.search(r"теоретические\s+положения", text, re.IGNORECASE)),
        "Критерии соответствия": bool(re.search(r"критерии\s+соответствия", text, re.IGNORECASE)),
        "Матрица соответствия": bool(re.search(r"матрица\s+соответствия", text, re.IGNORECASE)),
        "Методы проверки": bool(re.search(r"методы\s+проверки", text, re.IGNORECASE)),
        "Интеграционная проверка": bool(re.search(r"интеграционная\s+проверка", text, re.IGNORECASE))
    }
    return sections


def analyze_requirements(text, formulas):
    """Анализирует требования и их формализацию"""
    analysis = {
        "reliability": {
            "mentioned": bool(re.search(r"надежность", text, re.IGNORECASE)),
            "value": extract_value(text, r"надежность\s*[≥>]?\s*([\d.]+)%"),
            "formalized": any("надеж" in f.lower() for f in formulas)
        },
        "performance": {
            "mentioned": bool(re.search(r"производительность", text, re.IGNORECASE)),
            "value": extract_value(text, r"производительность\s*[≥>]?\s*([\d,]+)"),
            "formalized": any("производит" in f.lower() or "λ" in f for f in formulas)
        },
        "security": {
            "mentioned": bool(re.search(r"безопасность", text, re.IGNORECASE)),
            "value": extract_value(text, r"безопасность\s*[=:]\s*([\w\s]+)"),
            "formalized": any("безопас" in f.lower() for f in formulas)
        },
        "standards": re.findall(r'(ISO\s*\d+|IEEE\s*\d+|ГОСТ\s*\S+)', text)
    }
    return analysis


def extract_value(text, pattern):
    """Извлекает значение по шаблону"""
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(1) if match else "Не указано"


def analyze_verification(text):
    """Анализирует методы верификации"""
    analysis = {
        "automated_tests": {
            "mentioned": bool(re.search(r"автоматизированные\s+тесты", text, re.IGNORECASE)),
            "types": extract_list(text, r"(юнит-тесты|интеграционные\s+тесты|нагрузочное\s+тестирование)")
        },
        "expert_audit": {
            "mentioned": bool(re.search(r"экспертный\s+аудит", text, re.IGNORECASE)),
            "focus": extract_list(text, r"(анализ\s+кода|аудит\s+архитектуры|проверка\s+документации)")
        },
        "simulations": {
            "mentioned": bool(re.search(r"симуляции", text, re.IGNORECASE)),
            "scenarios": extract_list(text, r"(пиковая\s+нагрузка|отказ\s+оборудования|атака\s+безопасности)")
        }
    }
    return analysis


def extract_list(text, pattern):
    """Извлекает список значений по шаблону"""
    return list(set(re.findall(pattern, text, re.IGNORECASE)))


def generate_recommendations(sections, requirements, verification, formulas):
    """Генерирует детализированные рекомендации"""
    recommendations = []

    # Рекомендации по разделам
    for section, present in sections.items():
        if not present:
            recommendations.append(
                f"Добавьте раздел '{section}' с подробным описанием")

    # Рекомендации по требованиям
    if not formulas:
        recommendations.append(
            "Формализуйте требования с помощью математических моделей (пример: $R(t) \\geq 0.999$)")

    if not requirements["reliability"]["formalized"] and requirements["reliability"]["mentioned"]:
        recommendations.append(
            "Добавьте математическую формализацию для требования надежности (пример: $MTBF \\geq 1000$ часов)")

    if not requirements["performance"]["formalized"] and requirements["performance"]["mentioned"]:
        recommendations.append(
            "Формализуйте требования производительности (пример: $\\lambda \\geq 1000$ запросов/сек)")

    # Рекомендации по верификации
    if not verification["automated_tests"]["types"]:
        recommendations.append("Добавьте типы автоматизированных тестов: юнит - тесты, интеграционные ...

    if not verification["expert_audit"]["focus"]:
        recommendations.append(
            "Укажите фокусные области для экспертного аудита: анализ кода, аудит архитектуры")

    return recommendations

def save_report(report, original_path):
    """Сохраняет подробный отчет в текстовом формате"""
    report_path=os.path.splitext(original_path)[0] + "_ПолныйАнализ.txt"

    with open(report_path, 'w', encoding='utf-8') as f:
        # Заголовок
        f.write("ПОЛНЫЙ АНАЛИЗ ДОКУМЕНТА ПО МОДЕЛИ АДЕКВАТНОСТИ\n")
        f.write("=" * 100 + "\n\n")

        # Метаданные
        f.write(f"Документ: {report['metadata']['document']}\n")
        f.write(f"Дата анализа: {report['metadata']['analysis_date']}\n")
        f.write(f"Страниц: {report['metadata']['pages']}\n\n")

        # 1. Математические формулы
        f.write("1. МАТЕМАТИЧЕСКИЕ ФОРМУЛЫ\n")
        if report['formulas']:
            for i, formula in enumerate(report['formulas'], 1):
                f.write(f"   {i}. {formula}\n")
            f.write(f"   Всего формул: {len(report['formulas'])}\n")
        else:
            f.write("   Формулы не обнаружены\n")
        f.write("\n")

        # 2. Анализ разделов
        f.write("2. КЛЮЧЕВЫЕ РАЗДЕЛЫ\n")
        for section, present in report['sections'].items():
            status="✓ Присутствует" if present else "✗ Отсутствует"
            f.write(f"   - {section}: {status}\n")
        f.write("\n")

        # 3. Анализ требований
        f.write("3. ТРЕБОВАНИЯ И КРИТЕРИИ\n")
        req=report['requirements']
        f.write(f"   Надежность: {'✓' if req['reliability']['mentioned'] else '✗'} "
                f"({req['reliability']['value']}) "
                f"{'[Формализовано]' if req['reliability']['formalized'] else '[Требует формализации]'}\n")

        f.write(f"   Производительность: {'✓' if req['performance']['mentioned'] else '✗'} "
                f"({req['performance']['value']}) "
                f"{'[Формализовано]' if req['performance']['formalized'] else '[Требует формализации]'}\n")

        f.write(f"   Безопасность: {'✓' if req['security']['mentioned'] else '✗'} "
                f"({req['security']['value']}) "
                f"{'[Формализовано]' if req['security']['formalized'] else '[Требует формализации]'}\n")

        f.write(
            f"   Стандарты: {', '.join(req['standards']) if req['standards'] else 'Не указаны'}\n\n")

        # 4. Методы верификации
        f.write("4. МЕТОДЫ ВЕРИФИКАЦИИ\n")
        verif=report['verification']
        f.write("   Автоматизированные тесты:\n")
        f.write(
            f"      • Упомянуты: {'Да' if verif['automated_tests']['mentioned'] else 'Нет'}\n")
        f.write(f"      • Типы: {', '.join(verif['automated_tests']['types']) if verif['automated_te...

        f.write("   Экспертный аудит:\n")
        f.write(
            f"      • Упомянут: {'Да' if verif['expert_audit']['mentioned'] else 'Нет'}\n")
        f.write(f"      • Области: {', '.join(verif['expert_audit']['focus']) if verif['expert_audit...

        f.write("   Симуляции:\n")
        f.write(
            f"      • Упомянуты: {'Да' if verif['simulations']['mentioned'] else 'Нет'}\n")
        f.write(f"      • Сценарии: {', '.join(verif['simulations']['scenarios']) if verif['simulati...

        # 5. Рекомендации
        f.write("5. ДЕТАЛИЗИРОВАННЫЕ РЕКОМЕНДАЦИИ\n")
        if report['recommendations']:
            for i, rec in enumerate(report['recommendations'], 1):
                f.write(f"   {i}. {rec}\n")
        else:
            f.write(
                "   Документ полностью соответствует требованиям модели адекватности\n")

        # Заключение
        f.write("\n" + "=" * 100 + "\n")
        f.write("ЗАКЛЮЧЕНИЕ: ")
        issues = len(report['recommendations'])
        if issues == 0:
            f.write("Документ идеально соответствует модели адекватности систем")
        else:
            f.write(f"Обнаружено {issues} областей для улучшения")

    return report_path

if __name__ == "__main__":
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
    else:
        file_path = input("Перетащите файл .docx сюда: ").strip('"')

    result = analyze_document(file_path)

    if result.startswith("Ошибка"):
        printtttttttttttttttttttttttttttttttttt(result)
    else:
        printtttttttttttttttttttttttttttttttttt(
            f"Полный отчет сохранен: {result}")
        # Автоматически открываем отчет
        os.startfile(result)

    input("Нажмите Enter для выхода...")
