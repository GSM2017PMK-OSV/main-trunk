import os
import re
import sys
from collections import defaultdict
from datetime import datetime

from docx import Document


def deep_analyze_document(file_path):
    """Проводит глубокий анализ документа с извлечением формул и детальными рекомендациями"""
    try:
        doc = Document(file_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])

        # Извлечение ключевых элементов
        formulas = extract_formulas(full_text)
        requirements = analyze_requirements(full_text, formulas)
        decomposition = analyze_decomposition(full_text)
        verification = analyze_verification(full_text)
        matrix = analyze_compliance_matrix(doc.tables)
        integration = analyze_integration(full_text)

        # Формирование отчета
        report = {
            "metadata": {
                "document": os.path.basename(file_path),
                "analysis_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "pages": len(doc.paragraphs) // 50 + 1,
            },
            "formulas": formulas,
            "requirements": requirements,
            "decomposition": decomposition,
            "verification": verification,
            "matrix": matrix,
            "integration": integration,
            "recommendations": generate_recommendations(
                requirements, decomposition, verification, matrix, integration, formulas
            ),
        }

        # Сохранение полного отчета
        report_path = save_report(report, file_path)
        return report_path

    except Exception as e:
        return f"Ошибка анализа: {str(e)}"


def extract_formulas(text):
    """Извлекает и классифицирует математические формулы"""
    # Поиск формул в формате LaTeX
    latex_formulas = re.findall(r"\$(.*?)\$|\$\$(.*?)\$\$", text)
    formulas = [formula[0] or formula[1]
                for formula in latex_formulas if any(formula)]

    # Поиск формул в текстовом формате
    text_formulas = re.findall(
        r"Формула\s*[:\-]\s*(.+?)(?=\n|$)",
        text,
        re.IGNORECASE)

    # Классификация формул
    classified = defaultdict(list)
    for formula in formulas + text_formulas:
        formula = formula.strip()
        if not formula:
            continue

        # Определение типа формулы
        if re.search(r"R\(t\)|надежность", formula):
            classified["Надежность"].append(formula)
        elif re.search(r"\lambda|производительность", formula):
            classified["Производительность"].append(formula)
        elif re.search(r"V_i|функция\s*проверки", formula):
            classified["Функции проверки"].append(formula)
        elif re.search(r"M_{i,j}|матрица", formula):
            classified["Матрица соответствия"].append(formula)
        elif re.search(r"G\s*=\s*\(S,E\)|граф", formula):
            classified["Интеграция"].append(formula)
        else:
            classified["Другие"].append(formula)

    return dict(classified)


def analyze_requirements(text, formulas):
    """Анализ теоретических положений и требований"""
    analysis = {"found": [], "missing": [], "standards": [], "metrics": []}

    # Основные требования
    requirements = [
        "Надежность",
        "Производительность",
        "Безопасность",
        "Совместимость"]
    for req in requirements:
        if re.search(rf"{req}[\s\S]*?[\d≥>]", text, re.IGNORECASE):
            analysis["found"].append(req)
        else:
            analysis["missing"].append(req)

    # Стандарты
    standards = re.findall(r"(ISO\s*\d+|IEEE\s*\d+|ГОСТ\s*\S+)", text)
    analysis["standards"] = list(set(standards)) if standards else []

    # Измеримые параметры
    metrics = re.findall(r"([\w\s]+)\s*[=:>≥]\s*([\d\.,%]+)", text)
    analysis["metrics"] = [f"{m[0].strip()} = {m[1].strip()}" for m in metrics]

    # Проверка формализации требований
    analysis["formalized"] = bool(formulas.get(
        "Надежность") or formulas.get("Производительность"))

    return analysis


def analyze_decomposition(text):
    """Анализ структурной декомпозиции системы"""
    analysis = {
        "subsystems": {
            "management": check_completeness(text, "управляющие подсистемы", ["роли", "зоны ответственности"]),
            "programs": check_completeness(text, "программные компоненты", ["алгоритмы", "логика"]),
            "tech": check_completeness(text, "технологические решения", ["аппаратные средства", "протоколы"]),
        },
        "interaction": {
            "graph": bool(re.search(r"граф\s*взаимодействия", text, re.IGNORECASE)),
            "description": bool(re.search(r"описание\s*взаимодействия", text, re.IGNORECASE)),
        },
    }
    return analysis


def check_completeness(text, section, keywords):
    """Оценивает полноту описания раздела"""
    if not re.search(section, text, re.IGNORECASE):
        return "Отсутствует"

    coverage = sum(1 for kw in keywords if re.search(kw, text, re.IGNORECASE))
    return f"{coverage}/{len(keywords)} элементов" if coverage < len(
        keywords) else "Полное"


def analyze_verification(text):
    """Анализ методов проверки"""
    methods = {
        "automated": extract_test_methods(
            text, "автоматизированные тесты", [
                "юнит-тесты", "интеграционные", "нагрузочные"]
        ),
        "audit": extract_test_methods(
            text, "экспертный аудит", [
                "анализ кода", "аудит архитектуры", "проверка документации"]
        ),
        "simulations": extract_test_methods(
            text, "симуляции", [
                "моделирование нагрузки", "тестирование отказоустойчивости", "проверка безопасности"]
        ),
    }

    # Анализ инструментов
    tools = re.findall(r"(Selenium|JMeter|Postman|SonarQube|Checkmarx)", text)
    methods["tools"] = list(set(tools))

    return methods


def extract_test_methods(text, method_name, types):
    """Извлекает информацию о методах тестирования"""
    coverage = sum(1 for t in types if re.search(t, text, re.IGNORECASE))
    return {
        "mentioned": bool(re.search(method_name, text, re.IGNORECASE)),
        "coverage": f"{coverage}/{len(types)}",
        "missing": [t for t in types if not re.search(t, text, re.IGNORECASE)],
    }


def analyze_compliance_matrix(tables):
    """Анализ матрицы соответствия"""
    analysis = {"exists": False, "quality": "Низкая"}

    for table in tables:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        required_headers = ["теор", "подс", "метод", "рез", "док"]

        if any(req in header for header in headers for req in required_headers):
            analysis["exists"] = True

            # Оценка качества заполнения
            filled = 0
            for i, row in enumerate(table.rows):
                if i == 0:
                    continue  # Пропускаем заголовок
                if any(cell.text.strip() for cell in row.cells):
                    filled += 1

            completeness = filled / \
                (len(table.rows) - 1) if len(table.rows) > 1 else 0
            analysis["completeness"] = f"{completeness:.0%}"
            analysis["rows"] = len(table.rows) - 1

            # Оценка качества
            if completeness > 0.8:
                analysis["quality"] = "Высокая"
            elif completeness > 0.5:
                analysis["quality"] = "Средняя"
            break

    return analysis


def analyze_integration(text):
    """Анализ интеграционных проверок"""
    analysis = {
        "api": bool(re.search(r"проверка\s*API", text, re.IGNORECASE)),
        "data": bool(re.search(r"обмен\s*данными", text, re.IGNORECASE)),
        "protocols": re.findall(r"(HTTP(S)?/|gRPC|WebSocket|MQTT)", text),
        "scenarios": bool(re.search(r"сценарии\s*взаимодействия", text, re.IGNORECASE)),
    }
    return analysis


def generate_recommendations(req, decomp, verif, matrix, integ, formulas):
    """Генерирует детализированные практические рекомендации"""
    recommendations = []

    # Рекомендации по теоретическим основаниям
    if not formulas:
        recommendations.append(
            {
                "category": "Теоретические положения",
                "priority": "Высокая",
                "description": "Добавьте формальные математические модели для ключевых требований",
                "action": "Используйте LaTeX-нотацию для формул: $R(t) \\geq 0.999$",
            }
        )
    elif not formulas.get("Надежность"):
        recommendations.append(
            {
                "category": "Теоретические положения",
                "priority": "Средняя",
                "description": "Формализуйте требования надежности математически",
                "action": "Добавьте функцию надежности: $R(t) = e^{-\\lambda t}$",
            }
        )

    if not req["standards"]:
        recommendations.append(
            {
                "category": "Теоретические положения",
                "priority": "Высокая",
                "description": "Укажите стандарты соответствия",
                "action": "Добавьте ссылки на ISO, IEEE или ГОСТ стандарты",
            }
        )

    # Рекомендации по декомпозиции
    for subsystem, status in decomp["subsystems"].items():
        if "Отсутствует" in status:
            name = {
                "management": "Управляющих подсистем",
                "programs": "Программных компонентов",
                "tech": "Технологических решений",
            }[subsystem]

            recommendations.append(
                {
                    "category": "Структурная декомпозиция",
                    "priority": "Высокая",
                    "description": f"Добавьте описание {name}",
                    "action": f"Опишите роли, зоны ответственности и функции {name}",
                }
            )

    # Рекомендации по верификации
    for method, data in verif.items():
        if method == "tools":
            continue

        if not data["mentioned"]:
            name = {
                "automated": "Автоматизированных тестов",
                "audit": "Экспертного аудита",
                "simulations": "Симуляций",
            }[method]

            recommendations.append(
                {
                    "category": "Методы верификации",
                    "priority": "Высокая",
                    "description": f"Добавьте раздел о {name}",
                    "action": f"Опишите методы, инструменты и сценарии {name}",
                }
            )
        elif data["coverage"] != f"{len(data['missing'])+1}/3":
            for missing in data["missing"]:
                recommendations.append(
                    {
                        "category": "Методы верификации",
                        "priority": "Средняя",
                        "description": f"Добавьте описание {missing}",
                        "action": f"Подробно опишите процесс проведения {missing}",
                    }
                )

    # Рекомендации по матрице соответствия
    if not matrix["exists"]:
        recommendations.append(
            {
                "category": "Матрица соответствия",
                "priority": "Критическая",
                "description": "Добавьте матрицу соответствия",
                "action": "Создайте таблицу с колонками: Теорет. положение, Подсистема, Метод провер...
            }
        )
    elif matrix.get("quality", "Низкая") == "Низкая":
        recommendations.append(
            {
                "category": "Матрица соответствия",
                "priority": "Высокая",
                "description": f"Улучшите качество матрицы (текущее: {matrix.get('completeness', '0%')})",
                "action": "Заполните все пустые ячейки, добавьте ссылки на доказательства",
            }
        )

    # Рекомендации по интеграции
    if not integ["scenarios"]:
        recommendations.append(
            {
                "category": "Интеграция",
                "priority": "Высокая",
                "description": "Добавьте сценарии интеграционного тестирования",
                "action": "Опишите сценарии взаимодействия между подсистемами: нормальные условия, пиковая нагрузка, отказы",
            }
        )

    # Приоритизация рекомендаций
    priority_order = {
        "Критическая": 0,
        "Высокая": 1,
        "Средняя": 2,
        "Низкая": 3}
    return sorted(recommendations, key=lambda x: priority_order[x["priority"]])


def save_report(report, original_path):
    """Сохраняет детализированный отчет в текстовом формате"""
    report_path = os.path.splitext(original_path)[0] + "_ДетальныйАнализ.txt"

    with open(report_path, "w", encoding="utf-8") as f:
        # Заголовок отчета
        f.write("ДЕТАЛИЗИРОВАННЫЙ АНАЛИЗ АДЕКВАТНОСТИ СИСТЕМЫ\n")
        f.write("=" * 100 + "\n\n")

        # Основная информация
        meta = report["metadata"]
        f.write(f"Документ: {meta['document']}\n")
        f.write(f"Дата анализа: {meta['analysis_date']}\n")
        f.write(f"Страниц: {meta['pages']}\n\n")

        # 1. Формальные математические модели
        f.write("1. ФОРМАЛЬНЫЕ МАТЕМАТИЧЕСКИЕ МОДЕЛИ\n")
        if report["formulas"]:
            for category, formulas in report["formulas"].items():
                f.write(f"   {category}:\n")
                for formula in formulas:
                    f.write(f"      • {formula}\n")
            f.write(
                f"   Всего формул: {sum(len(f) for f in report['formulas'].values())}\n")
        else:
            f.write("   Математические модели не обнаружены\n")
        f.write("\n")

        # 2. Теоретические положения и требования
        f.write("2. ТЕОРЕТИЧЕСКИЕ ПОЛОЖЕНИЯ\n")
        req = report["requirements"]
        f.write(f"   Найдено требований: {', '.join(req['found'])}\n")
        if req["missing"]:
            f.write(
                f"   Отсутствующие требования: {', '.join(req['missing'])}\n")
        f.write(
            f"   Стандарты: {', '.join(req['standards']) or 'Не указаны'}\n")
        f.write(
            f"   Формализация: {'Присутствует' if req['formalized'] else 'Отсутствует'}\n")

        if req["metrics"]:
            f.write("   Измеримые параметры:\n")
            for metric in req["metrics"]:
                f.write(f"      • {metric}\n")
        else:
            f.write("   Измеримые параметры: Не обнаружены\n")
        f.write("\n")

        # 3. Структурная декомпозиция
        f.write("3. СТРУКТУРНАЯ ДЕКОМПОЗИЦИЯ СИСТЕМЫ\n")
        decomp = report["decomposition"]
        f.write(
            f"   Управляющие подсистемы: {decomp['subsystems']['management']}\n")
        f.write(
            f"   Программные компоненты: {decomp['subsystems']['programs']}\n")
        f.write(
            f"   Технологические решения: {decomp['subsystems']['tech']}\n")
        f.write(
            f"   Граф взаимодействия: {'Присутствует' if decomp['interaction']['graph'] else 'Отсутствует'}\n")
        f.write(
            f"   Описание взаимодействия: {'Присутствует' if decomp['interaction']['description'] else 'Отсутствует'}\n\n"
        )

        # 4. Методы верификации
        f.write("4. МЕТОДЫ ВЕРИФИКАЦИИ\n")
        verif = report["verification"]
        f.write("   Автоматизированные тесты:\n")
        f.write(
            f"      • Упомянуты: {'Да' if verif['automated']['mentioned'] else 'Нет'}\n")
        f.write(f"      • Покрытие: {verif['automated']['coverage']} типов\n")
        if verif["automated"]["missing"]:
            f.write(
                f"      • Отсутствует: {', '.join(verif['automated']['missing'])}\n")

        f.write("   Экспертный аудит:\n")
        f.write(
            f"      • Упомянут: {'Да' if verif['audit']['mentioned'] else 'Нет'}\n")
        f.write(f"      • Покрытие: {verif['audit']['coverage']} областей\n")
        if verif["audit"]["missing"]:
            f.write(
                f"      • Отсутствует: {', '.join(verif['audit']['missing'])}\n")

        f.write("   Симуляции:\n")
        f.write(
            f"      • Упомянуты: {'Да' if verif['simulations']['mentioned'] else 'Нет'}\n")
        f.write(
            f"      • Покрытие: {verif['simulations']['coverage']} сценариев\n")
        if verif["simulations"]["missing"]:
            f.write(
                f"      • Отсутствует: {', '.join(verif['simulations']['missing'])}\n")

        f.write(
            f"   Инструменты: {', '.join(verif['tools']) or 'Не указаны'}\n\n")

        # 5. Матрица соответствия
        f.write("5. МАТРИЦА СООТВЕТСТВИЯ\n")
        matrix = report["matrix"]
        f.write(f"   Присутствует: {'Да' if matrix['exists'] else 'Нет'}\n")
        if matrix["exists"]:
            f.write(f"   Полнота: {matrix.get('completeness', 'Н/Д')}\n")
            f.write(f"   Качество: {matrix['quality']}\n")
            f.write(f"   Строк: {matrix.get('rows', 0)}\n")
        f.write("\n")

        # 6. Интеграционные проверки
        f.write("6. ИНТЕГРАЦИОННЫЕ ПРОВЕРКИ\n")
        integ = report["integration"]
        f.write(f"   Проверка API: {'Да' if integ['api'] else 'Нет'}\n")
        f.write(f"   Обмен данными: {'Да' if integ['data'] else 'Нет'}\n")
        protocols = set(["".join(p) for p in integ["protocols"]])
        f.write(f"   Протоколы: {', '.join(protocols) or 'Не указаны'}\n")
        f.write(
            f"   Сценарии: {'Присутствуют' if integ['scenarios'] else 'Отсутствуют'}\n\n")

        # 7. Практические рекомендации
        f.write("7. ДЕТАЛИЗИРОВАННЫЕ РЕКОМЕНДАЦИИ\n")
        if report["recommendations"]:
            for i, rec in enumerate(report["recommendations"], 1):
                f.write(
                    f"   {i}. [{rec['priority']} приоритет] {rec['description']}\n")
                f.write(f"       Категория: {rec['category']}\n")
                f.write(f"       Действие: {rec['action']}\n")
                f.write("       " + "-" * 50 + "\n")
        else:
            f.write(
                "   Система полностью соответствует всем требованиям модели адекватности\n")

        # Заключение
        issues = len(report["recommendations"])
        f.write("\n" + "=" * 100 + "\n")
        f.write("ЗАКЛЮЧЕНИЕ: ")
        if issues == 0:
            f.write("Документ идеально соответствует модели адекватности систем")
        else:
            critical = sum(
                1 for r in report["recommendations"] if r["priority"] in [
                    "Критическая", "Высокая"])
            f.write(
                f"Обнаружено {issues} рекомендаций по улучшению, из них {critical} критически важных\n")

    return report_path


if __name__ == "__main__":
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
    else:
        file_path = input("Перетащите файл .docx сюда: ").strip('"')

    result = deep_analyze_document(file_path)

    if result.startswith("Ошибка"):
        printtttttttttttttttttttttttttttttt(result)
    else:
        printtttttttttttttttttttttttttttttt(
            f"Детализированный отчет сохранен: {result}")
        # Автоматически открываем отчет
        os.startfile(result)

    input("Нажмите Enter для выхода...")
