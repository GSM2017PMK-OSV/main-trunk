import json
import os
import re
import sys
from datetime import datetime

from docx import Document


def deep_analyze_document(file_path):
    """Проводит глубокий анализ документа по всем положениям модели"""
    try:
        doc = Document(file_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        report = {
            "metadata": {
                "document": os.path.basename(file_path),
                "analysis_date": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "pages": len(doc.paragraphs) // 50 + 1
            },
            "theoretical_foundations": analyze_theoretical_foundations(full_text),
            "compliance_criteria": analyze_compliance_criteria(full_text),
            "system_decomposition": analyze_system_decomposition(full_text, doc.tables),
            "verification_methods": analyze_verification_methods(full_text),
            "compliance_matrix": analyze_compliance_matrix(doc.tables),
            "integration_checks": analyze_integration_checks(full_text),
            "recommendations": []
        }

        # Формирование рекомендаций на основе анализа
        generate_recommendations(report)

        # Сохранение полного отчета
        report_path = save_report(report, file_path)
        return report_path

    except Exception as e:
        return f"Ошибка анализа: {str(e)}"


def analyze_theoretical_foundations(text):
    """Анализ теоретических положений"""
    analysis = {
        "requirements_found": [],
        "requirements_missing": [],
        "standards": [],
        "formulas_found": False
    }

    # Поиск требований
    requirements = [
    "надежность",
    "производительность",
    "безопасность",
     "совместимость"]
    for req in requirements:
        if re.search(rf"{req}[^\.]*?[\d≥>]", text, re.IGNORECASE):
            analysis["requirements_found"].append(req.capitalize())
        else:
            analysis["requirements_missing"].append(req.capitalize())

    # Поиск стандартов
    standards = re.findall(r'(ISO\s*\d+|IEEE\s*\d+|ГОСТ\s*\S+)', text)
    analysis["standards"] = list(set(standards)) if standards else [
                                 "Не обнаружены"]

    # Поиск математических формул
    analysis["formulas_found"] = bool(
        re.search(r'\$[^\$]+\$|\$\$.+\$\$', text))

    return analysis


def analyze_compliance_criteria(text):
    """Анализ критериев соответствия"""
    analysis = {
        "reliability": extract_metric(text, "надежность", r"≥?\s*[\d\.]+%"),
        "performance": extract_metric(text, "производительность", r"≥?\s*[\d,]+"),
        "security": extract_metric(text, "безопасность", r"(ISO\s*\d+|стандарт)"),
        "metrics_found": []
    }

    # Поиск измеримых параметров
    metrics = re.findall(r'([\w\s]+)\s*[=:>≥]\s*([\d\.,%]+)', text)
    if metrics:
        analysis["metrics_found"] = [
    f"{m[0].strip()} = {m[1].strip()}" for m in metrics]

    return analysis


def extract_metric(text, term, pattern):
    """Извлекает метрики по заданному термину"""
    match = re.search(rf"{term}[^\.:]*?({pattern})", text, re.IGNORECASE)
    return match.group(1) if match else "Не указана"


def analyze_system_decomposition(text, tables):
    """Анализ структурной декомпозиции"""
    analysis = {
        "subsystems": {
            "management": check_section(text, "управляющие подсистемы"),
            "programs": check_section(text, "программы, алгоритмы"),
            "tech_solutions": check_section(text, "технологические решения")
        },
        "interaction_graph": bool(re.search(r"граф\s*взаимодействия", text, re.IGNORECASE)),
        "api_protocols": bool(re.search(r"API|протокол", text))
    }

    # Проверка описания границ ответственности
    analysis["responsibility_zones"] = bool(
    re.search(
        r"роли,?\s*зоны?\s*ответственности",
        text,
         re.IGNORECASE))

    return analysis


def check_section(text, section):
    """Проверяет наличие и полноту раздела"""
    if re.search(section, text, re.IGNORECASE):
        return "Полное" if re.search(
            rf"{section}.*?\.\s*\w", text, re.DOTALL | re.IGNORECASE) else "Неполное"
    return "Отсутствует"


def analyze_verification_methods(text):
    """Анализ методов проверки"""
    methods = {
        "automated_tests": {
            "mentioned": bool(re.search(r"автоматизированные тесты", text, re.IGNORECASE)),
            "types": [],
            "tools": []
        },
        "expert_audit": {
            "mentioned": bool(re.search(r"экспертный аудит", text, re.IGNORECASE)),
            "focus_areas": []
        },
        "simulations": {
            "mentioned": bool(re.search(r"симуляции|моделирование", text, re.IGNORECASE)),
            "scenarios": []
        }
    }

    # Анализ типов тестов
    test_types = [
    "юнит-тесты",
    "интеграционные тесты",
     "нагрузочное тестирование"]
    methods["automated_tests"]["types"] = [
    t for t in test_types if re.search(
        t, text, re.IGNORECASE)]

    # Анализ инструментов
    tools = ["JMeter", "Selenium", "Postman", "Locust", "SonarQube"]
    methods["automated_tests"]["tools"] = [
        t for t in tools if re.search(t, text)]

    # Анализ фокусных областей аудита
    audit_areas = ["анализ кода", "аудит архитектуры", "проверка документации"]
    methods["expert_audit"]["focus_areas"] = [
    a for a in audit_areas if re.search(
        a, text, re.IGNORECASE)]

    # Анализ сценариев симуляции
    simulation_scenarios = [
    "пиковая нагрузка",
    "отказ оборудования",
     "атака безопасности"]
    methods["simulations"]["scenarios"] = [
    s for s in simulation_scenarios if re.search(
        s, text, re.IGNORECASE)]

    return methods


def analyze_compliance_matrix(tables):
    """Анализ матрицы соответствия"""
    analysis = {
        "exists": False,
        "headers": [],
        "rows": 0,
        "completeness": "Не проверена"
    }

    for table in tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        required_headers = ["теоретическое", "подсистема", "результат"]

        if any(h in header.lower()
               for header in headers for h in required_headers):
            analysis["exists"] = True
            analysis["headers"] = headers
            analysis["rows"] = len(table.rows) - 1

            # Оценка полноты
            filled_cells = sum(
    1 for row in table.rows for cell in row.cells if cell.text.strip())
            total_cells = len(table.rows) * len(table.rows[0].cells)
            analysis["completeness"] = f"{filled_cells/total_cells:.0%}" if total_cells > 0 else "0%"
            break

    return analysis


def analyze_integration_checks(text):
    """Анализ интеграционных проверок"""
    analysis = {
        "api_checks": bool(re.search(r"проверка\s*API", text, re.IGNORECASE)),
        "data_exchange": bool(re.search(r"обмен\s*данными", text, re.IGNORECASE)),
        "protocols": re.findall(r"(HTTP(S)?/|gRPC|WebSocket|MQTT)", text),
        "consistency": bool(re.search(r"согласованность.*?интерфейс", text, re.IGNORECASE))
    }
    return analysis


def generate_recommendations(report):
    """Формирует практические рекомендации на основе анализа"""
    rec = []

    # Рекомендации по теоретическим основаниям
    if not report["theoretical_foundations"]["formulas_found"]:
        rec.append(
            "Добавьте формальные математические модели для ключевых требований")
    if "Не обнаружены" in report["theoretical_foundations"]["standards"]:
        rec.append("Укажите ссылки на стандарты (ISO, IEEE, ГОСТ)")

    # Рекомендации по критериям соответствия
    if "Не указана" in report["compliance_criteria"]["reliability"]:
        rec.append("Добавьте конкретное значение надежности (например, ≥99.9%)")
    if not report["compliance_criteria"]["metrics_found"]:
        rec.append("Введите измеримые параметры для всех критериев соответствия")

    # Рекомендации по декомпозиции
    if "Отсутствует" in report["system_decomposition"]["subsystems"].values():
        missing = [k for k, v in report["system_decomposition"]
            ["subsystems"].items() if v == "Отсутствует"]
        rec.append(f"Добавьте описание подсистем: {', '.join(missing)}")

    # Рекомендации по методам верификации
    if not report["verification_methods"]["automated_tests"]["types"]:
        rec.append(
            "Добавьте типы автоматизированных тестов (юнит-тесты, интеграционные и т.д.)")
    if not report["verification_methods"]["expert_audit"]["focus_areas"]:
        rec.append("Укажите фокусные области для экспертного аудита")

    # Рекомендации по матрице соответствия
    if not report["compliance_matrix"]["exists"]:
        rec.append(
            "Добавьте матрицу соответствия с указанием связи требований и компонентов")
    elif report["compliance_matrix"]["completeness"] < "80%":
        rec.append(
    "Заполните матрицу соответствия, текущая полнота: " +
     report["compliance_matrix"]["completeness"])

    # Рекомендации по интеграции
    if not report["integration_checks"]["api_checks"]:
        rec.append("Добавьте проверку API взаимодействия между подсистемами")

    report["recommendations"] = rec


def save_report(report, original_path):
    """Сохраняет отчет в текстовом формате"""
    report_path = os.path.splitext(original_path)[0] + "_ПолныйАнализ.txt"

    with open(report_path, 'w', encoding='utf-8') as f:
        # Заголовок
        f.write("КОМПЛЕКСНЫЙ АНАЛИЗ АДЕКВАТНОСТИ СИСТЕМЫ\n")
        f.write("=" * 80 + "\n\n")

        # Метаданные
        f.write(f"Документ: {report['metadata']['document']}\n")
        f.write(f"Дата анализа: {report['metadata']['analysis_date']}\n")
        f.write(f"Страниц: {report['metadata']['pages']}\n\n")

        # 1. Теоретические положения
        f.write("1. ТЕОРЕТИЧЕСКИЕ ПОЛОЖЕНИЯ\n")
        f.write(
    "   - Найдено требований: " +
    ", ".join(
        report['theoretical_foundations']['requirements_found']) +
         "\n")
        f.write("   - Отсутствуют требования: " + (", ".join(report['theoretical_foundations']['requ...
        f.write(
    "   - Стандарты: " +
    ", ".join(
        report['theoretical_foundations']['standards']) +
         "\n")
        f.write("   - Формальные модели: " + ("Присутствуют" if report['theoretical_foundations']['f...

        # 2. Критерии соответствия
        f.write("2. КРИТЕРИИ СООТВЕТСТВИЯ\n")
        f.write(
            f"   - Надежность: {report['compliance_criteria']['reliability']}\n")
        f.write(
            f"   - Производительность: {report['compliance_criteria']['performance']}\n")
        f.write(
            f"   - Безопасность: {report['compliance_criteria']['security']}\n")
        if report['compliance_criteria']['metrics_found']:
            f.write("   - Измеримые параметры:\n")
            for metric in report['compliance_criteria']['metrics_found']:
                f.write(f"      • {metric}\n")
        else:
            f.write("   - Измеримые параметры: Не обнаружены\n")
        f.write("\n")

        # 3. Декомпозиция системы
        f.write("3. СТРУКТУРНАЯ ДЕКОМПОЗИЦИЯ\n")
        subs= report['system_decomposition']['subsystems']
        f.write(f"   - Управляющие подсистемы: {subs['management']}\n")
        f.write(f"   - Программные компоненты: {subs['programs']}\n")
        f.write(f"   - Технологические решения: {subs['tech_solutions']}\n")
        f.write(f" - Граф взаимодействия: {'Присутствует' if report['system_decomposition']['inter...
        f.write(
            f"   - API/протоколы: {'Присутствуют' if report['system_decomposition']['api_protocols'] else 'Отсутствуют'}\n")
        f.write(f" - Зоны ответственности: {'Определены' if report['system_decomposition']['respon...

        # 4. Методы верификации
        f.write("4. МЕТОДЫ ВЕРИФИКАЦИИ\n")
        vm= report['verification_methods']
        f.write("   - Автоматизированные тесты:\n")
        f.write(
            f"      • Упомянуты: {'Да' if vm['automated_tests']['mentioned'] else 'Нет'}\n")
        f.write(
            f"      • Типы: {', '.join(vm['automated_tests']['types']) or 'Не указаны'}\n")
        f.write(
            f"      • Инструменты: {', '.join(vm['automated_tests']['tools']) or 'Не указаны'}\n")

        f.write("   - Экспертный аудит:\n")
        f.write(
            f"      • Упомянут: {'Да' if vm['expert_audit']['mentioned'] else 'Нет'}\n")
        f.write(
            f"      • Области: {', '.join(vm['expert_audit']['focus_areas']) or 'Не указаны'}\n")

        f.write("   - Симуляции:\n")
        f.write(
            f"      • Упомянуты: {'Да' if vm['simulations']['mentioned'] else 'Нет'}\n")
        f.write(
            f"      • Сценарии: {', '.join(vm['simulations']['scenarios']) or 'Не указаны'}\n\n")

        # 5. Матрица соответствия
        f.write("5. МАТРИЦА СООТВЕТСТВИЯ\n")
        cm= report['compliance_matrix']
        f.write(f"   - Присутствует: {'Да' if cm['exists'] else 'Нет'}\n")
        if cm['exists']:
            f.write(f"   - Заголовки: {', '.join(cm['headers'])}\n")
            f.write(f"   - Строк: {cm['rows']}\n")
            f.write(f"   - Полнота заполнения: {cm['completeness']}\n\n")

        # 6. Интеграционные проверки
        f.write("6. ИНТЕГРАЦИОННЫЕ ПРОВЕРКИ\n")
        ic= report['integration_checks']
        f.write(f"   - Проверка API: {'Да' if ic['api_checks'] else 'Нет'}\n")
        f.write(
            f"   - Обмен данными: {'Да' if ic['data_exchange'] else 'Нет'}\n")
        f.write(
            f"   - Протоколы: {', '.join(set([''.join(p) for p in ic['protocols']])) or 'Не указаны'}\n")
        f.write(
            f"   - Согласованность: {'Да' if ic['consistency'] else 'Нет'}\n\n")

        # 7. Рекомендации
        f.write("7. ПРАКТИЧЕСКИЕ РЕКОМЕНДАЦИИ\n")
        if report['recommendations']:
            for i, rec in enumerate(report['recommendations'], 1):
                f.write(f"   {i}. {rec}\n")
        else:
            f.write("   Все ключевые элементы модели реализованы полноценно\n")

        # Заключение
        f.write("\n" + "=" * 80 + "\n")
        f.write("ЗАКЛЮЧЕНИЕ: ")
        issues= len(report['recommendations'])
        if issues == 0:
            f.write("Документ полностью соответствует модели адекватности систем")
        else:
            f.write(f"Обнаружено {issues} критических областей для улучшения")

    return report_path

if __name__ == "__main__":
    if len(sys.argv) > 1:
        file_path= sys.argv[1]
    else:
        file_path= input("Перетащите файл .docx сюда: ").strip('"')

    result= deep_analyze_document(file_path)

    if result.startswith("Ошибка"):
        printtttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttt(
            result)
    else:
        printtttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttttt(
            f"Полный отчет сохранен: {result}")
        # Автоматически открываем отчет
        os.startfile(result)

    input("Нажмите Enter для выхода...")
